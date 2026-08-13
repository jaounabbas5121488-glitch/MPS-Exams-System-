import json
import random
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

from database import get_db

def select_questions(conn, block_ids, mcq_count, short_count, long_count, comp_count):
    """Randomly select questions from given syllabus blocks."""
    # Fetch all questions for blocks
    blocks_str = ",".join(str(b) for b in block_ids)
    mcqs_pool = []
    short_pool = []
    long_pool = []
    comp_pool = []

    rows = conn.execute(f"""
        SELECT * FROM question_bank
        WHERE syllabus_block_id IN ({blocks_str})
        ORDER BY id DESC
    """).fetchall()

    for q in rows:
        if q["question_type"] == "mcq" and q["parent_comprehension_id"] is None:
            mcqs_pool.append(q)
        elif q["question_type"] == "short":
            short_pool.append(q)
        elif q["question_type"] == "long":
            long_pool.append(q)
        elif q["question_type"] == "comprehension":
            comp_pool.append(q)

    if len(mcqs_pool) < mcq_count:
        raise ValueError(f"Not enough MCQs in selected blocks. Available: {len(mcqs_pool)}, Required: {mcq_count}")
    if len(short_pool) < short_count:
        raise ValueError(f"Not enough Short Questions. Available: {len(short_pool)}, Required: {short_count}")
    if len(long_pool) < long_count:
        raise ValueError(f"Not enough Long Questions. Available: {len(long_pool)}, Required: {long_count}")
    if len(comp_pool) < comp_count:
        raise ValueError(f"Not enough Comprehensions. Available: {len(comp_pool)}, Required: {comp_count}")

    selected_mcqs = random.sample(mcqs_pool, mcq_count)
    selected_short = random.sample(short_pool, short_count)
    selected_long = random.sample(long_pool, long_count)
    selected_comp = random.sample(comp_pool, comp_count)

    # For each comprehension, fetch linked MCQs
    comp_with_mcqs = []
    for comp in selected_comp:
        linked_mcqs = conn.execute("""
            SELECT * FROM question_bank
            WHERE parent_comprehension_id = ?
        """, (comp["id"],)).fetchall()
        comp_with_mcqs.append({"passage": comp, "mcqs": linked_mcqs})

    return {
        "mcqs": selected_mcqs,
        "short": selected_short,
        "long": selected_long,
        "comprehensions": comp_with_mcqs,
    }

def generate_docx(selected, class_name, subject_name, blocks_names, total_marks,
                  mcq_marks, short_marks, long_marks, page_size="A4",
                  font_size=10, num_pages=1):
    doc = Document()
    # Set font and margins
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'  # or OMR for bubbles later
    font.size = Pt(font_size)
    section = doc.sections[0]
    if page_size == "A4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    elif page_size == "Legal":
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
    else:  # Letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Header lines
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mustafa Public School")
    run.bold = True
    run.font.size = Pt(16)

    info = f"Class: {class_name} | Subject: {subject_name} | Date: {date.today().strftime('%d-%m-%Y')} | Syllabus: {', '.join(blocks_names)} | Total Marks: {total_marks}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(info).font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Name: _______________ | Father Name: _______________ | Roll No: _______________")

    # Sections
    if selected["mcqs"]:
        doc.add_heading("MCQs", level=1)
        for idx, q in enumerate(selected["mcqs"], start=1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. {q['question_text']}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(f"A) ○   B) ○   C) ○   D) ○")
            # add options text
            p2 = doc.add_paragraph()
            p2.add_run(f"A) {q['option_a']}   B) {q['option_b']}   C) {q['option_c']}   D) {q['option_d']}")

    if selected["short"]:
        doc.add_heading("Short Questions", level=1)
        for idx, q in enumerate(selected["short"], start=1):
            doc.add_paragraph(f"{idx}. {q['question_text']}")

    if selected["long"]:
        doc.add_heading("Long Questions", level=1)
        for idx, q in enumerate(selected["long"], start=1):
            doc.add_paragraph(f"{idx}. {q['question_text']}")

    if selected["comprehensions"]:
        doc.add_heading("Comprehension", level=1)
        for comp in selected["comprehensions"]:
            doc.add_paragraph(comp["passage"]["comprehension_passage"])
            for idx, mcq in enumerate(comp["mcqs"], start=1):
                doc.add_paragraph(f"{idx}. {mcq['question_text']}")
                doc.add_paragraph(f"A) ○   B) ○   C) ○   D) ○")
                doc.add_paragraph(f"A) {mcq['option_a']}   B) {mcq['option_b']}   C) {mcq['option_c']}   D) {mcq['option_d']}")

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_xlsx(selected, class_name, subject_name, blocks_names, total_marks,
                  mcq_marks, short_marks, long_marks, page_size="A4",
                  font_size=10, num_pages=1):
    wb = Workbook()
    ws = wb.active
    ws.title = "Test"

    # Set font
    ws.sheet_view.rightToLeft = False  # later handle Urdu
    # Merge cells for school name
    ws.merge_cells('A1:J1')
    ws['A1'] = "Mustafa Public School"
    ws['A1'].font = Font(bold=True, size=16)
    ws['A1'].alignment = Alignment(horizontal='center')

    info = f"Class: {class_name} | Subject: {subject_name} | Date: {date.today().strftime('%d-%m-%Y')} | Syllabus: {', '.join(blocks_names)} | Total Marks: {total_marks}"
    ws.merge_cells('A2:J2')
    ws['A2'] = info
    ws['A2'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A3:J3')
    ws['A3'] = "Name: _______________ | Father Name: _______________ | Roll No: _______________"
    ws['A3'].alignment = Alignment(horizontal='center')

    row = 5
    # MCQs
    if selected["mcqs"]:
        ws.cell(row=row, column=1, value="MCQs").font = Font(bold=True)
        row += 1
        for idx, q in enumerate(selected["mcqs"], start=1):
            ws.cell(row=row, column=1, value=f"{idx}. {q['question_text']}")
            ws.cell(row=row, column=2, value=f"A) {q['option_a']}")
            ws.cell(row=row, column=3, value=f"B) {q['option_b']}")
            ws.cell(row=row, column=4, value=f"C) {q['option_c']}")
            ws.cell(row=row, column=5, value=f"D) {q['option_d']}")
            row += 1

    if selected["short"]:
        ws.cell(row=row, column=1, value="Short Questions").font = Font(bold=True)
        row += 1
        for idx, q in enumerate(selected["short"], start=1):
            ws.cell(row=row, column=1, value=f"{idx}. {q['question_text']}")
            row += 1

    if selected["long"]:
        ws.cell(row=row, column=1, value="Long Questions").font = Font(bold=True)
        row += 1
        for idx, q in enumerate(selected["long"], start=1):
            ws.cell(row=row, column=1, value=f"{idx}. {q['question_text']}")
            row += 1

    if selected["comprehensions"]:
        ws.cell(row=row, column=1, value="Comprehension").font = Font(bold=True)
        row += 1
        for comp in selected["comprehensions"]:
            ws.cell(row=row, column=1, value=comp["passage"]["comprehension_passage"])
            row += 1
            for idx, mcq in enumerate(comp["mcqs"], start=1):
                ws.cell(row=row, column=1, value=f"{idx}. {mcq['question_text']}")
                ws.cell(row=row, column=2, value=f"A) {mcq['option_a']}")
                ws.cell(row=row, column=3, value=f"B) {mcq['option_b']}")
                ws.cell(row=row, column=4, value=f"C) {mcq['option_c']}")
                ws.cell(row=row, column=5, value=f"D) {mcq['option_d']}")
                row += 1

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def generate_answer_key_docx(selected, class_name, subject_name, blocks_names, total_marks):
    doc = Document()
    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mustafa Public School - Answer Key")
    run.bold = True
    run.font.size = Pt(14)

    info = f"Class: {class_name} | Subject: {subject_name} | Date: {date.today().strftime('%d-%m-%Y')} | Syllabus: {', '.join(blocks_names)}"
    doc.add_paragraph(info)

    # MCQs answers
    if selected["mcqs"]:
        doc.add_heading("MCQs", level=1)
        for idx, q in enumerate(selected["mcqs"], start=1):
            doc.add_paragraph(f"{idx}. Correct: {q['correct_answer']}")

    if selected["comprehensions"]:
        doc.add_heading("Comprehension MCQs", level=1)
        for comp in selected["comprehensions"]:
            for idx, mcq in enumerate(comp["mcqs"], start=1):
                doc.add_paragraph(f"{idx}. Correct: {mcq['correct_answer']}")

    # Short/Long answer if available
    if selected["short"]:
        doc.add_heading("Short Questions (Answers if provided)", level=1)
        for idx, q in enumerate(selected["short"], start=1):
            if q["answer_text"]:
                doc.add_paragraph(f"{idx}. {q['answer_text']}")
            else:
                doc.add_paragraph(f"{idx}. [No answer stored]")

    if selected["long"]:
        doc.add_heading("Long Questions (Answers if provided)", level=1)
        for idx, q in enumerate(selected["long"], start=1):
            if q["answer_text"]:
                doc.add_paragraph(f"{idx}. {q['answer_text']}")
            else:
                doc.add_paragraph(f"{idx}. [No answer stored]")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output