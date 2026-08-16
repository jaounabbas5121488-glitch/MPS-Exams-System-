import io
import re
import os
import html as html_lib
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment


def safe_str(value):
    return "" if value is None else str(value)


def strip_html(raw_html):
    """Remove HTML tags, unescape entities, and return plain text."""
    text = re.sub(r'<[^>]+>', '', safe_str(raw_html))
    text = html_lib.unescape(text)
    return text.strip()


def is_rtl_text(text):
    """Return True if text contains Arabic/Urdu/Persian characters."""
    rtl_chars = re.search(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]', text)
    return bool(rtl_chars)


def extract_images_from_html(raw_html):
    """Return list of image src paths from <img> tags."""
    raw_html = safe_str(raw_html)
    if not raw_html:
        return []
    pattern = r'<img[^>]+src="([^"]+)"'
    return re.findall(pattern, raw_html, re.IGNORECASE)


def resolve_image_path(src):
    """Convert image src to a valid filesystem path."""
    src = src.replace('\\', '/')
    src = src.lstrip('/')
    while src.startswith('../'):
        src = src[3:]
    candidates = [
        src,
        f"uploads/{src}",
        f"/{src}",
        f"uploads/question_images/{os.path.basename(src)}",
        f"uploads/{os.path.basename(src)}",
        os.path.join("uploads", "question_images", os.path.basename(src)),
    ]
    for candidate in candidates:
        candidate_path = os.path.normpath(candidate)
        if os.path.exists(candidate_path):
            return candidate_path
    return None


def set_paragraph_rtl(paragraph, rtl=True):
    """Set paragraph direction to RTL or LTR in the underlying XML."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove existing bidi element
    for bidi in pPr.findall(qn('w:bidi')):
        pPr.remove(bidi)
    if rtl:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        # Insert bidi in correct position (before jc, after spacing/ind etc.)
        # For simplicity, insert before jc if exists, else append.
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            jc.addprevious(bidi)
        else:
            pPr.append(bidi)


def set_run_rtl(run, rtl=True):
    """Set run properties for RTL/LTR."""
    rPr = run._r.get_or_add_rPr()
    for rtl_elem in rPr.findall(qn('w:rtl')):
        rPr.remove(rtl_elem)
    if rtl:
        rtl_elem = OxmlElement('w:rtl')
        rtl_elem.set(qn('w:val'), '1')
        rPr.append(rtl_elem)


def determine_direction(text, default_direction):
    """Determine paragraph direction based on subject default and text content."""
    if default_direction == 'rtl':
        return 'rtl'
    elif default_direction == 'ltr':
        return 'ltr'
    else:  # 'both' or 'auto'
        return 'rtl' if is_rtl_text(text) else 'ltr'


def add_inline_content(paragraph, raw_html, default_direction='ltr'):
    """
    Add text and images inline to the given paragraph.
    Sets RTL/LTR alignment and direction on both paragraph and runs.
    """
    text = strip_html(raw_html)
    direction = determine_direction(text, default_direction)
    rtl = (direction == 'rtl')
    if rtl:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(paragraph, True)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_rtl(paragraph, False)

    if text:
        run = paragraph.add_run(text)
        set_run_rtl(run, rtl)
    for img_src in extract_images_from_html(raw_html):
        img_path = resolve_image_path(img_src)
        if img_path and os.path.exists(img_path):
            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(2))
            set_run_rtl(run, rtl)
        else:
            run = paragraph.add_run(f" [Image missing: {img_src}] ")
            set_run_rtl(run, rtl)


def add_mcq_options(doc, mcq_data, default_direction='ltr'):
    """
    Adds MCQ options with OMR-style bubbles.
    If total options text is short, one line; otherwise each option on separate line.
    """
    bubble_letters = ['Ⓐ', 'Ⓑ', 'Ⓒ', 'Ⓓ']
    opt_texts = []
    for i, opt_key in enumerate(['option_a', 'option_b', 'option_c', 'option_d']):
        opt_text = strip_html(mcq_data[opt_key])
        opt_texts.append(f"{bubble_letters[i]} {opt_text}")

    total_len = sum(len(t) for t in opt_texts)
    if total_len <= 80:  # Fit on one line
        p = doc.add_paragraph()
        combined_text = "     ".join(opt_texts)
        direction = determine_direction(combined_text, default_direction)
        rtl = (direction == 'rtl')
        if rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_rtl(p, True)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_rtl(p, False)
        run = p.add_run(combined_text)
        set_run_rtl(run, rtl)
    else:
        # One option per line
        for opt_text in opt_texts:
            p = doc.add_paragraph()
            direction = determine_direction(opt_text, default_direction)
            rtl = (direction == 'rtl')
            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(p, True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_rtl(p, False)
            run = p.add_run(opt_text)
            set_run_rtl(run, rtl)


def generate_rich_docx(selected, class_name, subject_name, blocks_names, total_marks,
                       mcq_marks, short_marks, long_marks, page_size="A4",
                       font_size=10, num_pages=1, default_direction='ltr'):
    doc = Document()
    section = doc.sections[0]
    if page_size == "A4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    elif page_size == "Legal":
        section.page_width = Inches(8.5)
        section.page_height = Inches(14)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Header (always centered)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mustafa Public School")
    run.bold = True
    run.font.size = Pt(16)

    info = f"Class: {class_name} | Subject: {subject_name} | Date: {date.today().strftime('%d-%m-%Y')} | Syllabus: {', '.join(blocks_names)} | Total Marks: {total_marks}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(info).font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Name: _______________ | Father Name: _______________ | Roll No: _______________")

    # Helper to add heading with direction
    def add_heading(text, default_direction):
        heading = doc.add_heading(text, level=1)
        direction = determine_direction(text, default_direction)
        rtl = (direction == 'rtl')
        if rtl:
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_rtl(heading, True)
        else:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_rtl(heading, False)
        return heading

    # MCQs
    if selected["mcqs"]:
        add_heading("MCQs", default_direction)
        for idx, q in enumerate(selected["mcqs"], start=1):
            p = doc.add_paragraph()
            # Determine direction based on question text
            q_text = strip_html(q['question_text'])
            direction = determine_direction(q_text, default_direction)
            rtl = (direction == 'rtl')
            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(p, True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_rtl(p, False)
            # Add number run with direction
            num_run = p.add_run(f"{idx}. ")
            set_run_rtl(num_run, rtl)
            # Add question content
            add_inline_content(p, q['question_text'], default_direction)
            # Add options
            add_mcq_options(doc, q, default_direction)

    # Short Questions
    if selected["short"]:
        add_heading("Short Questions", default_direction)
        for idx, q in enumerate(selected["short"], start=1):
            p = doc.add_paragraph()
            q_text = strip_html(q['question_text'])
            direction = determine_direction(q_text, default_direction)
            rtl = (direction == 'rtl')
            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(p, True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_rtl(p, False)
            num_run = p.add_run(f"{idx}. ")
            set_run_rtl(num_run, rtl)
            add_inline_content(p, q['question_text'], default_direction)

    # Long Questions
    if selected["long"]:
        add_heading("Long Questions", default_direction)
        for idx, q in enumerate(selected["long"], start=1):
            p = doc.add_paragraph()
            q_text = strip_html(q['question_text'])
            direction = determine_direction(q_text, default_direction)
            rtl = (direction == 'rtl')
            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(p, True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_rtl(p, False)
            num_run = p.add_run(f"{idx}. ")
            set_run_rtl(num_run, rtl)
            add_inline_content(p, q['question_text'], default_direction)

    # Comprehensions
    if selected["comprehensions"]:
        add_heading("Comprehension", default_direction)
        for comp in selected["comprehensions"]:
            p = doc.add_paragraph()
            passage_text = strip_html(comp['passage']['comprehension_passage'])
            direction = determine_direction(passage_text, default_direction)
            rtl = (direction == 'rtl')
            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_paragraph_rtl(p, True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_rtl(p, False)
            add_inline_content(p, comp['passage']['comprehension_passage'], default_direction)
            for idx, mcq in enumerate(comp['mcqs'], start=1):
                p = doc.add_paragraph()
                mcq_text = strip_html(mcq['question_text'])
                direction = determine_direction(mcq_text, default_direction)
                rtl = (direction == 'rtl')
                if rtl:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    set_paragraph_rtl(p, True)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_rtl(p, False)
                num_run = p.add_run(f"{idx}. ")
                set_run_rtl(num_run, rtl)
                add_inline_content(p, mcq['question_text'], default_direction)
                add_mcq_options(doc, mcq, default_direction)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_rich_xlsx(selected, class_name, subject_name, blocks_names, total_marks,
                       mcq_marks, short_marks, long_marks, page_size="A4",
                       font_size=10, num_pages=1, default_direction='ltr'):
    wb = Workbook()
    ws = wb.active
    ws.title = "Test"
    if default_direction == 'rtl':
        ws.sheet_view.rightToLeft = True
    else:
        ws.sheet_view.rightToLeft = False

    ws.merge_cells('A1:J1')
    ws['A1'] = "Mustafa Public School"
    ws['A1'].font = Font(bold=True, size=16)
    ws['A1'].alignment = Alignment(horizontal='center')

    info = f"Class: {class_name} | Subject: {subject_name} | Date: {date.today().strftime('%d-%m-%Y')} | Syllabus: {', '.join(blocks_names)} | Total Marks: {total_marks}"
    ws.merge_cells('A2:J2')
    ws['A2'] = info
    ws['A2'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A3:J3')
    ws['A3'] = "Name: _______________ | Father Name: _______________ | Roll No: _______________"
    ws['A3'].alignment = Alignment(horizontal='center')

    row = 5
    if selected["mcqs"]:
        ws.cell(row=row, column=1, value="MCQs").font = Font(bold=True)
        row += 1
        for idx, q in enumerate(selected["mcqs"], start=1):
            text = strip_html(q['question_text'])
            ws.cell(row=row, column=1, value=f"{idx}. {text}")
            row += 1
            opt_texts = []
            bubble_letters = ['Ⓐ', 'Ⓑ', 'Ⓒ', 'Ⓓ']
            for i, opt_key in enumerate(['option_a', 'option_b', 'option_c', 'option_d']):
                opt_text = strip_html(q[opt_key])
                opt_texts.append(f"{bubble_letters[i]} {opt_text}")
            ws.cell(row=row, column=1, value="     ".join(opt_texts))
            row += 1

    if selected["short"]:
        ws.cell(row=row, column=1, value="Short Questions").font = Font(bold=True)
        row += 1
        for idx, q in enumerate(selected["short"], start=1):
            text = strip_html(q['question_text'])
            ws.cell(row=row, column=1, value=f"{idx}. {text}")
            row += 1

    if selected["long"]:
        ws.cell(row=row, column=1, value="Long Questions").font = Font(bold=True)
        row += 1
        for idx, q in enumerate(selected["long"], start=1):
            text = strip_html(q['question_text'])
            ws.cell(row=row, column=1, value=f"{idx}. {text}")
            row += 1

    if selected["comprehensions"]:
        ws.cell(row=row, column=1, value="Comprehension").font = Font(bold=True)
        row += 1
        for comp in selected["comprehensions"]:
            passage_text = strip_html(comp['passage']['comprehension_passage'])
            ws.cell(row=row, column=1, value=passage_text)
            row += 1
            for idx, mcq in enumerate(comp['mcqs'], start=1):
                text = strip_html(mcq['question_text'])
                ws.cell(row=row, column=1, value=f"{idx}. {text}")
                row += 1
                opt_texts = []
                bubble_letters = ['Ⓐ', 'Ⓑ', 'Ⓒ', 'Ⓓ']
                for i, opt_key in enumerate(['option_a', 'option_b', 'option_c', 'option_d']):
                    opt_text = strip_html(mcq[opt_key])
                    opt_texts.append(f"{bubble_letters[i]} {opt_text}")
                ws.cell(row=row, column=1, value="     ".join(opt_texts))
                row += 1

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def generate_answer_key_rich_docx(selected, class_name, subject_name, blocks_names, total_marks):
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mustafa Public School - Answer Key")
    run.bold = True
    run.font.size = Pt(14)

    info = f"Class: {class_name} | Subject: {subject_name} | Date: {date.today().strftime('%d-%m-%Y')} | Syllabus: {', '.join(blocks_names)}"
    doc.add_paragraph(info)

    if selected["mcqs"]:
        doc.add_heading("MCQs", level=1)
        for idx, q in enumerate(selected["mcqs"], start=1):
            doc.add_paragraph(f"{idx}. Correct: {q['correct_answer']}")

    if selected["comprehensions"]:
        doc.add_heading("Comprehension MCQs", level=1)
        for comp in selected["comprehensions"]:
            for idx, mcq in enumerate(comp['mcqs'], start=1):
                doc.add_paragraph(f"{idx}. Correct: {mcq['correct_answer']}")

    if selected["short"]:
        doc.add_heading("Short Questions (Answers if provided)", level=1)
        for idx, q in enumerate(selected["short"], start=1):
            if q["answer_text"]:
                doc.add_paragraph(f"{idx}. {strip_html(q['answer_text'])}")
            else:
                doc.add_paragraph(f"{idx}. [No answer stored]")

    if selected["long"]:
        doc.add_heading("Long Questions (Answers if provided)", level=1)
        for idx, q in enumerate(selected["long"], start=1):
            if q["answer_text"]:
                doc.add_paragraph(f"{idx}. {strip_html(q['answer_text'])}")
            else:
                doc.add_paragraph(f"{idx}. [No answer stored]")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output